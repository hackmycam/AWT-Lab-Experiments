"""
Script to take screenshots of all AWT experiments and create a Word document.
"""
import os
import time
import subprocess
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import UnexpectedAlertPresentException, NoAlertPresentException
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = r"c:\Users\praveen R\Desktop\AWT"
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "screenshots")
os.makedirs(SCREENSHOTS_DIR, exist_ok=True)

def get_driver():
    options = Options()
    options.add_argument("--headless=new")
    options.add_argument("--window-size=1280,900")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--allow-file-access-from-files")
    driver = webdriver.Chrome(options=options)
    return driver

def dismiss_alerts(driver):
    """Dismiss any open alert dialogs."""
    try:
        alert = driver.switch_to.alert
        alert.accept()
        time.sleep(0.5)
    except NoAlertPresentException:
        pass

def take_screenshot(driver, url, filename, wait_time=3):
    filepath = os.path.join(SCREENSHOTS_DIR, filename)
    dismiss_alerts(driver)
    try:
        driver.get(url)
    except UnexpectedAlertPresentException:
        dismiss_alerts(driver)
        driver.get(url)
    time.sleep(wait_time)
    dismiss_alerts(driver)
    try:
        driver.save_screenshot(filepath)
    except UnexpectedAlertPresentException:
        dismiss_alerts(driver)
        driver.save_screenshot(filepath)
    print(f"  Screenshot saved: {filename}")
    return filepath

def start_process(cmd, cwd):
    """Start a background process and return it."""
    proc = subprocess.Popen(
        cmd, cwd=cwd, shell=True,
        stdout=subprocess.PIPE, stderr=subprocess.PIPE
    )
    return proc

def kill_process(proc):
    """Kill a process safely."""
    try:
        proc.terminate()
        proc.wait(timeout=5)
    except:
        try:
            proc.kill()
        except:
            pass

def kill_port(port):
    """Kill whatever is running on a port."""
    try:
        result = subprocess.run(
            f'netstat -ano | findstr ":{port} "',
            shell=True, capture_output=True, text=True
        )
        for line in result.stdout.strip().split('\n'):
            if 'LISTENING' in line:
                pid = line.strip().split()[-1]
                subprocess.run(f'taskkill /F /PID {pid}', shell=True, capture_output=True)
    except:
        pass

def main():
    driver = get_driver()
    screenshots = {}
    processes = []

    # ===== Experiment 1: University Website (HTML frames) =====
    print("Experiment 1: University Website")
    exp1_path = os.path.join(BASE_DIR, "experiment-1", "index.html").replace("\\", "/")
    screenshots["exp1"] = take_screenshot(driver, f"file:///{exp1_path}", "exp1_output.png")

    # ===== Experiment 2: Form Validation =====
    print("Experiment 2: Form Validation")
    exp2_path = os.path.join(BASE_DIR, "experiment-2", "form.html").replace("\\", "/")
    screenshots["exp2"] = take_screenshot(driver, f"file:///{exp2_path}", "exp2_output.png")

    # ===== Experiment 3: React Employee/Events Demo =====
    print("Experiment 3: React Employee Demo")
    kill_port(5175)
    exp3_proc = start_process(["npm", "run", "dev", "--", "--port", "5175"], os.path.join(BASE_DIR, "experiment-3"))
    time.sleep(8)
    screenshots["exp3"] = take_screenshot(driver, "http://localhost:5175", "exp3_output.png")
    kill_process(exp3_proc)
    kill_port(5175)

    # ===== Experiment 4: React Lifecycle Demo =====
    print("Experiment 4: React Lifecycle Demo")
    kill_port(5176)
    exp4_proc = start_process(["npm", "run", "dev", "--", "--port", "5176"], os.path.join(BASE_DIR, "experiment-4"))
    time.sleep(8)
    screenshots["exp4"] = take_screenshot(driver, "http://localhost:5176", "exp4_output.png")
    kill_process(exp4_proc)
    kill_port(5176)

    # ===== Experiment 5: React App =====
    print("Experiment 5: React App")
    kill_port(5177)
    exp5_proc = start_process(["npm", "run", "dev", "--", "--port", "5177"], os.path.join(BASE_DIR, "experiment-5"))
    time.sleep(8)
    screenshots["exp5"] = take_screenshot(driver, "http://localhost:5177", "exp5_output.png")
    kill_process(exp5_proc)
    kill_port(5177)

    # ===== Experiment 6: React Router =====
    print("Experiment 6: React Router")
    kill_port(5178)
    exp6_proc = start_process(["npm", "run", "dev", "--", "--port", "5178"], os.path.join(BASE_DIR, "experiment-6"))
    time.sleep(8)
    screenshots["exp6"] = take_screenshot(driver, "http://localhost:5178", "exp6_output.png")
    kill_process(exp6_proc)
    kill_port(5178)

    # ===== Experiment 7: Library API =====
    print("Experiment 7: Library API")
    kill_port(3000)
    time.sleep(1)
    exp7_proc = start_process(["node", "server.js"], os.path.join(BASE_DIR, "experiment-7-library-api"))
    time.sleep(3)
    screenshots["exp7"] = take_screenshot(driver, "http://localhost:3000/books", "exp7_output.png")
    kill_process(exp7_proc)
    kill_port(3000)
    time.sleep(1)

    # ===== Experiment 8: MongoDB CRUD (terminal output) =====
    print("Experiment 8: MongoDB CRUD")
    exp8_result = subprocess.run(
        ["node", "app.js"],
        cwd=os.path.join(BASE_DIR, "experiment-8-mongodb-crud"),
        capture_output=True, text=True, timeout=15
    )
    exp8_output = exp8_result.stdout
    print(f"  Exp 8 output captured ({len(exp8_output)} chars)")

    # ===== Experiment 9: Login System =====
    print("Experiment 9: Login System")
    kill_port(5000)
    kill_port(5174)
    time.sleep(1)
    
    # Start backend
    exp9_backend = start_process(["node", "server.js"], os.path.join(BASE_DIR, "experiment-9", "backend"))
    time.sleep(3)
    
    # Start frontend
    exp9_frontend = start_process(["npm", "run", "dev", "--", "--port", "5174"], os.path.join(BASE_DIR, "experiment-9", "frontend"))
    time.sleep(8)
    
    screenshots["exp9"] = take_screenshot(driver, "http://localhost:5174", "exp9_form.png")
    
    kill_process(exp9_frontend)
    kill_process(exp9_backend)
    kill_port(5174)
    kill_port(5000)

    # ===== Experiment 10: Person App =====
    print("Experiment 10: Person App")
    kill_port(3000)
    time.sleep(1)
    exp10_proc = start_process(["node", "server.js"], os.path.join(BASE_DIR, "experiment-10-person-app"))
    time.sleep(3)
    
    screenshots["exp10_form"] = take_screenshot(driver, "http://localhost:3000", "exp10_form.png")
    screenshots["exp10_data"] = take_screenshot(driver, "http://localhost:3000/show", "exp10_data.png")
    
    # Keep exp 10 running for Compass viewing
    processes.append(exp10_proc)

    driver.quit()

    # ========== Create Word Document ==========
    print("\nCreating Word Document...")
    doc = Document()

    # Title
    title = doc.add_heading("AWT Lab Experiments - Output Screenshots", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    experiments = [
        {
            "num": 1,
            "title": "Experiment 1: Design a Website Using HTML with Frames",
            "desc": "A university website using HTML frames to display home, faculty, courses, and menu pages.",
            "images": [("exp1", "University Website Output")]
        },
        {
            "num": 2,
            "title": "Experiment 2: Form Validation Using JavaScript",
            "desc": "An HTML form with JavaScript validation for name, email, phone, and password fields.",
            "images": [("exp2", "Form Validation Output")]
        },
        {
            "num": 3,
            "title": "Experiment 3: React Components - Employee and Events Demo",
            "desc": "A React application demonstrating multiple components with employee data and event handling.",
            "images": [("exp3", "React Employee Component Output")]
        },
        {
            "num": 4,
            "title": "Experiment 4: React Component Lifecycle Methods",
            "desc": "A React application demonstrating component lifecycle methods (componentDidMount, componentDidUpdate, componentWillUnmount).",
            "images": [("exp4", "React Lifecycle Demo Output")]
        },
        {
            "num": 5,
            "title": "Experiment 5: React State Management",
            "desc": "A React application demonstrating state management concepts.",
            "images": [("exp5", "React State Management Output")]
        },
        {
            "num": 6,
            "title": "Experiment 6: React Router Navigation",
            "desc": "A React application demonstrating client-side routing with React Router.",
            "images": [("exp6", "React Router Output")]
        },
        {
            "num": 7,
            "title": "Experiment 7: RESTful API with Express.js",
            "desc": "A REST API for library book management using Express.js with GET, POST, PUT, DELETE operations.",
            "images": [("exp7", "API Response - GET /books")]
        },
        {
            "num": 8,
            "title": "Experiment 8: MongoDB CRUD Operations with Node.js",
            "desc": "Perform CRUD operations (Create, Read, Update, Delete) with Node.js and MongoDB using Mongoose.",
            "terminal_output": exp8_output,
            "images": []
        },
        {
            "num": 9,
            "title": "Experiment 9: Login Credentials Using ReactJS, NodeJS, MongoDB",
            "desc": "Create Login credentials using ReactJS and NodeJS. Upon successful login display 'Login Success' otherwise 'Login Failed' using MongoDB.",
            "images": [("exp9", "Login Form")]
        },
        {
            "num": 10,
            "title": "Experiment 10: Person Details Using NodeJS, Express, MongoDB",
            "desc": "A NodeJS application to insert and display details of a person using HTML, Express framework, and MongoDB.",
            "images": [("exp10_form", "Person Details Form"), ("exp10_data", "Stored Data (/show)")]
        }
    ]

    for exp in experiments:
        doc.add_page_break()
        doc.add_heading(exp["title"], level=1)
        doc.add_paragraph(exp["desc"])
        doc.add_heading("Output:", level=2)

        # Terminal output for exp 8
        if "terminal_output" in exp and exp["terminal_output"]:
            p = doc.add_paragraph()
            run = p.add_run("Terminal Output:")
            run.bold = True
            run.font.size = Pt(11)
            
            terminal_para = doc.add_paragraph()
            terminal_run = terminal_para.add_run(exp["terminal_output"])
            terminal_run.font.name = "Consolas"
            terminal_run.font.size = Pt(9)
            terminal_run.font.color.rgb = RGBColor(0, 100, 0)

        # Screenshots
        for img_key, caption in exp["images"]:
            if img_key in screenshots and os.path.exists(screenshots[img_key]):
                doc.add_paragraph(caption, style="Caption")
                doc.add_picture(screenshots[img_key], width=Inches(6))
            else:
                doc.add_paragraph(f"[Screenshot not available: {caption}]")

        # Note for MongoDB experiments
        if exp["num"] in [8, 9, 10]:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("MongoDB Compass: Connect to mongodb://127.0.0.1:27017 to view the data.")
            run.bold = True
            run.italic = True
            run.font.size = Pt(10)
            db_info_map = {8: ("college", "students"), 9: ("loginDB", "users"), 10: ("persondb", "people")}
            db_name, col_name = db_info_map[exp["num"]]
            db_para = doc.add_paragraph()
            db_run = db_para.add_run(f"Database: {db_name} -> Collection: {col_name}")
            db_run.font.size = Pt(10)

    # Save document
    output_path = os.path.join(BASE_DIR, "AWT_Lab_Experiments_Output.docx")
    doc.save(output_path)
    print(f"\nWord document saved to: {output_path}")
    
    # Cleanup remaining processes
    for proc in processes:
        kill_process(proc)
    
    print("Done! You can now add MongoDB Compass screenshots manually.")

if __name__ == "__main__":
    main()
